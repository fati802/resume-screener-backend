"""
DOCX parser — extracts raw text from Word document resume files.
"""

from docx import Document


async def parse_docx(file_path: str) -> str:
    """
    Extract all text from a DOCX file.
    Returns the full text as a single string.
    """
    doc = Document(file_path)
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text.strip())

    return "\n".join(text_parts)